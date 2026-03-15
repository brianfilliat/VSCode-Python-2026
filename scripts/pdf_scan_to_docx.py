#!/usr/bin/env python3
"""
Convert a scanned (image-based) PDF to a Microsoft Word .docx document
using EasyOCR for text recognition.

Usage:
    python scripts/pdf_scan_to_docx.py --pdf Theft-Network-scan3-11-2026.pdf \
        --out outputs/Theft-Network-scan3-11-2026.docx

Requires (already in venv):
    pymupdf, easyocr, python-docx, torch
"""

import argparse
import sys
import os

# Force UTF-8 stdout/stderr so EasyOCR's download progress bar
# (uses █ U+2588) doesn't crash on Windows cp1252 consoles.
if sys.stdout.encoding and sys.stdout.encoding.lower() != "utf-8":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

import fitz  # pymupdf
import easyocr
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def pdf_pages_as_images(pdf_path: str, dpi: int = 200):
    """Yield (page_number, numpy_array) for each page in the PDF."""
    import numpy as np
    doc = fitz.open(pdf_path)
    mat = fitz.Matrix(dpi / 72, dpi / 72)  # scale factor
    for i, page in enumerate(doc):
        pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(
            pix.height, pix.width, 3
        )
        yield i + 1, img
    doc.close()


def ocr_image(reader: easyocr.Reader, img):
    """Run OCR on a numpy image array, return extracted text lines."""
    results = reader.readtext(img, detail=0, paragraph=True)
    return results


def build_docx(pages_text: list[tuple[int, list[str]]], out_path: str):
    doc = Document()

    # --- Title page ---
    title_para = doc.add_heading("Theft Network — Transcribed Notes", level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Source: Theft-Network-scan3-11-2026.pdf")
    doc.add_paragraph(f"Total pages: {len(pages_text)}")
    doc.add_page_break()

    for page_num, lines in pages_text:
        # Page heading
        heading = doc.add_heading(f"Page {page_num}", level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if lines:
            for line in lines:
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(2)
        else:
            p = doc.add_paragraph("[No text detected on this page]")
            p.runs[0].italic = True

        doc.add_paragraph("")  # spacer between pages

    doc.save(out_path)
    print(f"Saved: {out_path}")


def main():
    parser = argparse.ArgumentParser(
        description="Convert scanned PDF to Word document using EasyOCR"
    )
    parser.add_argument("--pdf", required=True, help="Path to scanned PDF")
    parser.add_argument("--out", required=True, help="Output .docx path")
    parser.add_argument(
        "--dpi",
        type=int,
        default=200,
        help="Render DPI for page images (higher = better accuracy, slower)",
    )
    parser.add_argument(
        "--lang",
        nargs="+",
        default=["en"],
        help="EasyOCR language codes (default: en)",
    )
    parser.add_argument(
        "--gpu",
        action="store_true",
        default=False,
        help="Use GPU acceleration if available",
    )
    args = parser.parse_args()

    if not os.path.isfile(args.pdf):
        print(f"Error: PDF not found: {args.pdf}", file=sys.stderr)
        sys.exit(1)

    out_dir = os.path.dirname(os.path.abspath(args.out))
    os.makedirs(out_dir, exist_ok=True)

    print(f"Loading EasyOCR reader (languages: {args.lang}, gpu={args.gpu}) ...")
    reader = easyocr.Reader(args.lang, gpu=args.gpu)

    pages_text = []
    total_pages = fitz.open(args.pdf).page_count
    print(f"Processing {total_pages} pages from: {args.pdf}")

    for page_num, img in pdf_pages_as_images(args.pdf, dpi=args.dpi):
        print(f"  OCR page {page_num}/{total_pages} ...", end="\r", flush=True)
        lines = ocr_image(reader, img)
        pages_text.append((page_num, lines))

    print(f"\nBuilding Word document ...")
    build_docx(pages_text, args.out)
    print("Done.")


if __name__ == "__main__":
    main()
